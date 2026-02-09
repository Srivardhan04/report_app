"""
Report generator – produces PDF (WeasyPrint) and DOCX (python-docx) reports.
Color theme: KL University red/black/white.
Attendance % cell is color-coded (no separate Status column).
"""

import os
import base64
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import List

# Set DLL path for WeasyPrint BEFORE importing it
os.environ.setdefault("WEASYPRINT_DLL_DIRECTORIES", r"C:\msys64\mingw64\bin")

from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import (
    BASE_DIR, LOGO_PATH, UNIVERSITY_NAME, UNIVERSITY_FULL_NAME,
    DEPARTMENT_NAME, HOD_NAME, REPORT_TITLE, REPORT_DIR, ENABLE_TELUGU_NOTICE,
)
from models.student import StudentProfile


# ──────────────────────────────────────────────
#  Jinja2 Setup
# ──────────────────────────────────────────────
_template_env = Environment(
    loader=FileSystemLoader(str(BASE_DIR / "templates")),
    autoescape=True,
)


def _get_logo_base64() -> str:
    """Read KL logo and return as base64 data URI for embedding in HTML."""
    if LOGO_PATH.exists():
        with open(LOGO_PATH, "rb") as f:
            data = base64.b64encode(f.read()).decode()
        return f"data:image/png;base64,{data}"
    return ""


# ══════════════════════════════════════════════
#  PDF Generation (WeasyPrint)
# ══════════════════════════════════════════════

def generate_pdf(student: StudentProfile) -> bytes:
    """Generate a single-page A4 PDF report for one student."""
    template = _template_env.get_template("report.html")
    html_content = template.render(
        student=student,
        logo_base64=_get_logo_base64(),
        university_name=UNIVERSITY_NAME,
        university_full_name=UNIVERSITY_FULL_NAME,
        department_name=DEPARTMENT_NAME,
        report_title=REPORT_TITLE,
        hod_name=HOD_NAME,
        enable_telugu_notice=ENABLE_TELUGU_NOTICE,
        generated_date=datetime.now().strftime("%B %d, %Y"),
    )
    pdf_bytes = HTML(string=html_content).write_pdf()
    return pdf_bytes


def save_pdf(student: StudentProfile) -> Path:
    """Generate and save PDF to disk. Returns file path."""
    pdf_bytes = generate_pdf(student)
    filename = f"{student.student_id}_{student.student_name.replace(' ', '_')}_report.pdf"
    path = REPORT_DIR / filename
    path.write_bytes(pdf_bytes)
    return path


# ══════════════════════════════════════════════
#  Bulk ZIP Generation
# ══════════════════════════════════════════════

def generate_all_pdfs_zip(students: List[StudentProfile]) -> bytes:
    """Generate PDF for every student and bundle into a ZIP archive."""
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for student in students:
            pdf_bytes = generate_pdf(student)
            filename = f"{student.student_id}_{student.student_name.replace(' ', '_')}_report.pdf"
            zf.writestr(filename, pdf_bytes)
    buffer.seek(0)
    return buffer.read()


def save_all_pdfs_zip(students: List[StudentProfile]) -> Path:
    """Generate ZIP and save to disk."""
    zip_bytes = generate_all_pdfs_zip(students)
    path = REPORT_DIR / "all_student_reports.zip"
    path.write_bytes(zip_bytes)
    return path


# ══════════════════════════════════════════════
#  DOCX Generation (python-docx)
# ══════════════════════════════════════════════

# --- KL color constants for DOCX ---
_KL_RED = RGBColor(0xC6, 0x28, 0x28)
_KL_DARK = RGBColor(0x22, 0x22, 0x22)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GRAY = RGBColor(0x55, 0x55, 0x55)


def _add_centered_text(doc: Document, text: str, size: int = 12,
                       bold: bool = False, color: RGBColor = None,
                       space_after: int = 2):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return para


def _set_cell_shading(cell, hex_color: str):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    shading.set(qn("w:val"), "clear")


def _att_docx_color(attendance_class: str):
    """Return (text_color, bg_hex) for DOCX attendance cell."""
    return {
        "att-red":    (RGBColor(0xB7, 0x1C, 0x1C), "FDECEA"),
        "att-yellow": (RGBColor(0xF5, 0x7F, 0x17), "FFF8E1"),
        "att-green":  (RGBColor(0x1B, 0x5E, 0x20), "E8F5E9"),
    }.get(attendance_class, (RGBColor(0x22, 0x22, 0x22), "FFFFFF"))


def generate_docx(student: StudentProfile) -> bytes:
    """Generate a DOCX report for a single student (KL red theme)."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # --- Logo ---
    if LOGO_PATH.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.2))

    # --- University branding ---
    _add_centered_text(doc, "KL University", size=18, bold=True, color=_KL_RED)
    _add_centered_text(doc, DEPARTMENT_NAME, size=11, bold=True, color=_KL_DARK)
    _add_centered_text(doc, REPORT_TITLE, size=13, bold=True, color=_KL_RED, space_after=4)
    _add_centered_text(
        doc, f"Date: {datetime.now().strftime('%B %d, %Y')}",
        size=9, color=_GRAY, space_after=8
    )

    # ═══ Student Details ═══
    doc.add_paragraph().add_run("Student Details").bold = True
    details = [
        ("Student ID", student.student_id),
        ("Student Name", student.student_name),
    ]
    if student.section:
        details.append(("Section", student.section))
    if student.year:
        details.append(("Year", student.year))
    if student.semester:
        details.append(("Semester", student.semester))
    if student.branch:
        details.append(("Branch", student.branch))

    table = doc.add_table(rows=len(details), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value) if value else "N/A"
        if table.rows[i].cells[0].paragraphs[0].runs:
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ═══ Attendance Table (NO Status column — color in Attendance % cell) ═══
    if student.attendance_records:
        doc.add_paragraph().add_run("Current Semester Attendance").bold = True

        att_table = doc.add_table(rows=1, cols=5)
        att_table.style = "Light Grid Accent 1"
        att_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Subject Code", "Subject Name", "Classes Held",
                    "Classes Attended", "Attendance %"]
        for j, h in enumerate(headers):
            cell = att_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_shading(cell, "C62828")
            cell.paragraphs[0].runs[0].font.color.rgb = _WHITE

        for att in student.attendance_records:
            row = att_table.add_row()
            row.cells[0].text = att.subject_code
            row.cells[1].text = att.subject_name
            row.cells[2].text = str(att.classes_held)
            row.cells[3].text = str(att.classes_attended)
            row.cells[4].text = f"{att.attendance_percentage:.1f}%"

            # Color-code the attendance % cell
            txt_color, bg_hex = _att_docx_color(att.attendance_class)
            _set_cell_shading(row.cells[4], bg_hex)
            if row.cells[4].paragraphs[0].runs:
                row.cells[4].paragraphs[0].runs[0].font.color.rgb = txt_color
                row.cells[4].paragraphs[0].runs[0].bold = True

        # Overall attendance
        para = doc.add_paragraph()
        run = para.add_run(f"Overall Attendance: {student.overall_attendance:.1f}%")
        run.bold = True
        if student.overall_attendance < 75:
            run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        elif student.overall_attendance < 80:
            run.font.color.rgb = RGBColor(0xF5, 0x7F, 0x17)
        else:
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

        # Attendance warning (immediately after overall attendance)
        low_att = [a for a in student.attendance_records if a.attendance_percentage < 75.0]
        if low_att:
            subj_list = ", ".join(
                f"{a.subject_name} ({a.attendance_percentage:.1f}%)" for a in low_att
            )
            warning_en = (
                f"Notice Regarding Low Attendance: "
                f"It is hereby brought to the notice of the parent/guardian that the student's "
                f"attendance is below the minimum required threshold of 75% in the following "
                f"subject(s): {subj_list}. As per KL University regulations, students who fail "
                f"to maintain 75% attendance are liable for detention and may be debarred from "
                f"end-semester examinations."
            )
            p_warn = doc.add_paragraph()
            p_warn.paragraph_format.space_before = Pt(4)
            run_w = p_warn.add_run(warning_en)
            run_w.font.size = Pt(8.5)
            run_w.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)

            if ENABLE_TELUGU_NOTICE:
                warning_te = (
                    "సూచన: విద్యార్థి హాజరు శాతం పై పేర్కొన్న సబ్జెక్టు(ల)లో 75% కంటే తక్కువగా "
                    "ఉన్నది. కేఎల్ విశ్వవిద్యాలయ నిబంధనల ప్రకారం 75% హాజరును నిర్వహించని "
                    "విద్యార్థులు డిటెన్షన్‌కు గురవుతారు."
                )
                p_te = doc.add_paragraph()
                run_te = p_te.add_run(warning_te)
                run_te.font.size = Pt(8)
                run_te.italic = True
                run_te.font.color.rgb = _GRAY

        doc.add_paragraph()

    # ═══ Previous Semester Results ═══
    if student.previous_results:
        doc.add_paragraph().add_run("Previous Semester Results").bold = True

        res_table = doc.add_table(rows=1, cols=4)
        res_table.style = "Light Grid Accent 1"
        res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Subject Code", "Subject Name", "Grade", "Credits"]
        for j, h in enumerate(headers):
            cell = res_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_shading(cell, "C62828")
            cell.paragraphs[0].runs[0].font.color.rgb = _WHITE

        for res in student.previous_results:
            row = res_table.add_row()
            row.cells[0].text = res.subject_code
            row.cells[1].text = res.subject_name
            row.cells[2].text = res.grade
            row.cells[3].text = str(res.credits)
            if res.is_backlog:
                _set_cell_shading(row.cells[2], "FDECEA")
                if row.cells[2].paragraphs[0].runs:
                    row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)

        # Previous Semester CGPA
        if student.cgpa is not None:
            para = doc.add_paragraph()
            run = para.add_run(f"Previous Semester CGPA: {student.cgpa:.2f}")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)  # deep indigo

        if student.backlog_count > 0:
            backlog_subjs = ', '.join(student.backlog_subjects)
            backlog_en = (
                f"Notice Regarding Backlogs: "
                f"It is observed that the student has backlog(s) in the following subject(s): "
                f"{backlog_subjs}. The student is advised to take dedicated academic effort to "
                f"clear the above-mentioned backlog(s) at the earliest. The department will "
                f"provide necessary academic guidance and closely monitor the student's progress."
            )
            p_bl = doc.add_paragraph()
            p_bl.paragraph_format.space_before = Pt(4)
            run_bl = p_bl.add_run(backlog_en)
            run_bl.font.size = Pt(8.5)
            run_bl.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)

            if ENABLE_TELUGU_NOTICE:
                backlog_te = (
                    f"సూచన: విద్యార్థికి క్రింది సబ్జెక్టు(ల)లో — {backlog_subjs} — "
                    "బ్యాక్\u200cలాగ్(లు) ఉన్నట్లు గమనించబడింది. విద్యార్థి వీలైనంత త్వరగా "
                    "క్లియర్ చేయడానికి అంకితమైన అకడమిక్ కృషి చేయాలని సూచించబడింది."
                )
                p_bl_te = doc.add_paragraph()
                run_bl_te = p_bl_te.add_run(backlog_te)
                run_bl_te.font.size = Pt(8)
                run_bl_te.italic = True
                run_bl_te.font.color.rgb = _GRAY
        else:
            para = doc.add_paragraph()
            run = para.add_run("No Backlogs — All Subjects Cleared")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

        doc.add_paragraph()

    # ═══ Counselor Details ═══
    if student.counselor_name or student.counselor_id or student.counselor_email or student.counselor_phone:
        doc.add_paragraph().add_run("Counselor / Mentor Details").bold = True
        c_details = [
            ("Counselor Name", student.counselor_name or "N/A"),
            ("Counselor ID", student.counselor_id or "N/A"),
            ("Counselor Email", student.counselor_email or "N/A"),
            ("Counselor Phone", student.counselor_phone or "N/A"),
        ]
        c_table = doc.add_table(rows=len(c_details), cols=2)
        c_table.style = "Light Grid Accent 1"
        for i, (label, value) in enumerate(c_details):
            c_table.rows[i].cells[0].text = label
            c_table.rows[i].cells[1].text = value
            if c_table.rows[i].cells[0].paragraphs[0].runs:
                c_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

    # ═══ Footer (sign-off only — no duplicate notices) ═══
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)

    run = para.add_run("Sincerely,\n")
    run.bold = True
    run.font.size = Pt(11)

    run = para.add_run("Head of the Department\n")
    run.bold = True
    run.font.size = Pt(11)

    run = para.add_run(f"{HOD_NAME}\n")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = _KL_RED

    run = para.add_run(DEPARTMENT_NAME)
    run.font.size = Pt(10)
    run.font.color.rgb = _GRAY

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def save_docx(student: StudentProfile) -> Path:
    docx_bytes = generate_docx(student)
    filename = f"{student.student_id}_{student.student_name.replace(' ', '_')}_report.docx"
    path = REPORT_DIR / filename
    path.write_bytes(docx_bytes)
    return path
